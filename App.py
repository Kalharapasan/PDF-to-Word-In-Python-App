import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
import fitz  # PyMuPDF library
from docx import Document
from docx.shared import Inches
import subprocess
import platform

class PDFToWordConverter:
    def __init__(self, master):
        self.master = master
        master.title("PDF to Word Converter Pro")
        master.geometry("600x650")
        master.configure(bg='#2c3e50')
        master.resizable(False, False)
        
        # Center the window
        self.center_window()

        # Style configuration
        self.setup_styles()

        # Variables
        self.selected_files = []
        self.is_converting = False

        # Create and set up GUI components
        self.create_widgets()

    def center_window(self):
        self.master.update_idletasks()
        x = (self.master.winfo_screenwidth() // 2) - (600 // 2)
        y = (self.master.winfo_screenheight() // 2) - (650 // 2)
        self.master.geometry(f"600x650+{x}+{y}")

    def setup_styles(self):
        # Color scheme
        self.bg_primary = '#2c3e50'
        self.bg_secondary = '#34495e'
        self.accent_color = '#3498db'
        self.success_color = '#27ae60'
        self.error_color = '#e74c3c'
        self.text_primary = '#ecf0f1'
        self.text_secondary = '#bdc3c7'
        
        # Configure ttk styles
        style = ttk.Style()
        style.theme_use('clam')
        
        # Button style
        style.configure('Custom.TButton',
                       background=self.accent_color,
                       foreground='white',
                       font=('Segoe UI', 10, 'bold'),
                       relief='flat',
                       borderwidth=0)
        
        style.map('Custom.TButton',
                 background=[('active', '#2980b9')])
        
        # Progress bar style
        style.configure('Custom.Horizontal.TProgressbar',
                       background=self.accent_color,
                       troughcolor=self.bg_secondary,
                       borderwidth=0,
                       lightcolor=self.accent_color,
                       darkcolor=self.accent_color)

    def create_widgets(self):
        # Main container
        main_frame = tk.Frame(self.master, bg=self.bg_primary)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Title
        title_label = tk.Label(
            main_frame,
            text="PDF to Word Converter Pro",
            font=('Segoe UI', 18, 'bold'),
            bg=self.bg_primary,
            fg=self.text_primary
        )
        title_label.pack(pady=(0, 30))

        # File selection frame
        file_frame = tk.Frame(main_frame, bg=self.bg_secondary, relief='raised', bd=2)
        file_frame.pack(fill=tk.X, pady=(0, 20))

        file_title = tk.Label(
            file_frame,
            text="📄 File Selection",
            font=('Segoe UI', 12, 'bold'),
            bg=self.bg_secondary,
            fg=self.text_primary
        )
        file_title.pack(pady=(15, 10))

        # Button frame
        button_frame = tk.Frame(file_frame, bg=self.bg_secondary)
        button_frame.pack(pady=(0, 10))

        self.select_single_btn = ttk.Button(
            button_frame,
            text="Select Single PDF",
            style='Custom.TButton',
            command=self.select_single_file
        )
        self.select_single_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.select_multiple_btn = ttk.Button(
            button_frame,
            text="Select Multiple PDFs",
            style='Custom.TButton',
            command=self.select_multiple_files
        )
        self.select_multiple_btn.pack(side=tk.LEFT)

        # Selected files display
        self.files_listbox = tk.Listbox(
            file_frame,
            height=4,
            bg=self.bg_primary,
            fg=self.text_primary,
            selectbackground=self.accent_color,
            font=('Consolas', 9),
            relief='flat',
            borderwidth=0
        )
        self.files_listbox.pack(fill=tk.X, padx=15, pady=(10, 10))

        # Clear files button
        self.clear_btn = ttk.Button(
            file_frame,
            text="Clear All",
            command=self.clear_files
        )
        self.clear_btn.pack(pady=(5, 10))

        # Conversion options frame
        options_frame = tk.Frame(main_frame, bg=self.bg_secondary, relief='raised', bd=2)
        options_frame.pack(fill=tk.X, pady=(0, 20))

        options_title = tk.Label(
            options_frame,
            text="⚙️ Conversion Options",
            font=('Segoe UI', 12, 'bold'),
            bg=self.bg_secondary,
            fg=self.text_primary
        )
        options_title.pack(pady=(15, 10))

        # Checkboxes for options
        self.preserve_formatting = tk.BooleanVar(value=True)
        self.open_output_folder = tk.BooleanVar(value=True)

        preserve_check = tk.Checkbutton(
            options_frame,
            text="Preserve basic formatting",
            variable=self.preserve_formatting,
            bg=self.bg_secondary,
            fg=self.text_primary,
            selectcolor=self.bg_primary,
            activebackground=self.bg_secondary,
            activeforeground=self.text_primary,
            font=('Segoe UI', 10)
        )
        preserve_check.pack(pady=5)

        folder_check = tk.Checkbutton(
            options_frame,
            text="Open output folder after conversion",
            variable=self.open_output_folder,
            bg=self.bg_secondary,
            fg=self.text_primary,
            selectcolor=self.bg_primary,
            activebackground=self.bg_secondary,
            activeforeground=self.text_primary,
            font=('Segoe UI', 10)
        )
        folder_check.pack(pady=(5, 10))

        # Convert button - Make it more prominent and ensure it's visible
        convert_frame = tk.Frame(main_frame, bg=self.bg_primary)
        convert_frame.pack(fill=tk.X, pady=(15, 15))
        
        self.convert_button = tk.Button(
            convert_frame,
            text="🔄 Convert to Word",
            font=('Segoe UI', 14, 'bold'),
            bg=self.accent_color,
            fg='white',
            activebackground='#2980b9',
            activeforeground='white',
            relief='flat',
            borderwidth=0,
            pady=15,
            command=self.start_conversion,
            cursor='hand2'
        )
        self.convert_button.pack(fill=tk.X, padx=30)

        # Progress frame
        progress_frame = tk.Frame(main_frame, bg=self.bg_primary)
        progress_frame.pack(fill=tk.X, pady=(0, 20))

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            style='Custom.Horizontal.TProgressbar'
        )
        self.progress_bar.pack(fill=tk.X, pady=(0, 10))

        # Status label
        self.status_label = tk.Label(
            main_frame,
            text="Ready to convert PDF files",
            font=('Segoe UI', 10),
            bg=self.bg_primary,
            fg=self.text_secondary,
            wraplength=550
        )
        self.status_label.pack()

        # Update button state
        self.update_convert_button_state()

    def select_single_file(self):
        file_path = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF files", "*.pdf")]
        )
        
        if file_path:
            self.selected_files = [file_path]
            self.update_files_display()

    def select_multiple_files(self):
        file_paths = filedialog.askopenfilenames(
            title="Select PDF Files",
            filetypes=[("PDF files", "*.pdf")]
        )
        
        if file_paths:
            self.selected_files = list(file_paths)
            self.update_files_display()

    def clear_files(self):
        self.selected_files = []
        self.update_files_display()

    def update_files_display(self):
        self.files_listbox.delete(0, tk.END)
        for file_path in self.selected_files:
            filename = os.path.basename(file_path)
            self.files_listbox.insert(tk.END, filename)
        
        self.update_convert_button_state()

    def update_convert_button_state(self):
        if self.selected_files and not self.is_converting:
            self.convert_button.config(
                state=tk.NORMAL,
                bg=self.accent_color,
                text="🔄 Convert to Word"
            )
        else:
            self.convert_button.config(
                state=tk.DISABLED,
                bg='#7f8c8d',
                text="Select PDF files first" if not self.selected_files else "Converting..."
            )

    def update_status(self, message, color=None):
        if color is None:
            color = self.text_secondary
        self.status_label.config(text=message, fg=color)
        self.master.update_idletasks()

    def start_conversion(self):
        if not self.selected_files:
            messagebox.showerror("Error", "Please select at least one PDF file.")
            return

        self.is_converting = True
        self.update_convert_button_state()
        self.progress_var.set(0)
        
        # Start conversion in a separate thread
        thread = threading.Thread(target=self.convert_files)
        thread.daemon = True
        thread.start()

    def convert_files(self):
        successful_conversions = []
        failed_conversions = []
        total_files = len(self.selected_files)
        
        for i, pdf_path in enumerate(self.selected_files):
            try:
                # Update progress
                progress = (i / total_files) * 100
                self.progress_var.set(progress)
                
                filename = os.path.basename(pdf_path)
                self.update_status(f"Converting: {filename}")
                
                # Generate output path
                docx_path = os.path.splitext(pdf_path)[0] + ".docx"
                
                # Convert PDF to Word
                self.convert_single_pdf(pdf_path, docx_path)
                
                successful_conversions.append(os.path.basename(docx_path))
                
            except Exception as e:
                failed_conversions.append(f"{os.path.basename(pdf_path)}: {str(e)}")

        # Complete progress
        self.progress_var.set(100)
        
        # Show results
        self.show_conversion_results(successful_conversions, failed_conversions)
        
        # Reset state
        self.is_converting = False
        self.update_convert_button_state()

    def convert_single_pdf(self, pdf_path, docx_path):
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        
        # Create Word document
        doc = Document()
        
        # Set document margins for better formatting
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Extract text from each page
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Get text blocks with position information
            text_blocks = page.get_text("blocks")
            
            if page_num > 0:  # Add page break for subsequent pages
                doc.add_page_break()
            
            # Process each text block
            for block in text_blocks:
                if len(block) >= 5:  # Text block has position info
                    block_text = block[4].strip()
                    
                    if block_text:
                        # Split by lines to preserve line breaks
                        lines = block_text.split('\n')
                        
                        for line in lines:
                            line = line.strip()
                            if line:
                                # Check if it's a header/title (short line, likely bold)
                                if len(line) < 100 and any(keyword in line.lower() for keyword in 
                                    ['lab', 'sheet', 'exercise', 'practical', 'assignment']):
                                    # Add as heading
                                    heading = doc.add_heading(line, level=1)
                                    heading.alignment = 0  # Left align
                                
                                # Check if it's a subheading or exercise number
                                elif line.startswith(('01.', '02.', '03.', '04.', '05.', '06.', '07.', '08.', '09.', '10.')) or \
                                     line.lower().startswith(('exercise', 'question')):
                                    # Add as subheading
                                    subheading = doc.add_heading(line, level=2)
                                    subheading.alignment = 0  # Left align
                                
                                # Check if it's code or list content (indented or contains special characters)
                                elif any(char in line for char in ['[', ']', '=', '(', ')', '"']) or \
                                     line.startswith(('    ', '\t')) or \
                                     any(word in line.lower() for word in ['color_list', 'sample list', 'lists:']):
                                    # Add as code/monospace paragraph
                                    p = doc.add_paragraph()
                                    run = p.add_run(line)
                                    run.font.name = 'Courier New'
                                    run.font.size = Inches(0.12)  # 12pt
                                    p.paragraph_format.left_indent = Inches(0.5)  # Indent code blocks
                                    p.paragraph_format.space_after = Inches(0.1)
                                
                                # Regular paragraph
                                else:
                                    p = doc.add_paragraph(line)
                                    p.paragraph_format.space_after = Inches(0.1)
                        
                        # Add extra spacing after each block
                        if len(lines) > 1:
                            doc.add_paragraph()

        # If no content was added, fall back to simple text extraction
        if len(doc.paragraphs) == 0:
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    if page_num > 0:
                        doc.add_page_break()
                    
                    # Process line by line
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)

        # Save document
        doc.save(docx_path)
        pdf_document.close()

    def show_conversion_results(self, successful, failed):
        result_message = ""
        
        if successful:
            result_message += f"✅ Successfully converted {len(successful)} file(s):\n"
            for file in successful[:5]:  # Show first 5
                result_message += f"  • {file}\n"
            if len(successful) > 5:
                result_message += f"  ... and {len(successful) - 5} more\n"
        
        if failed:
            result_message += f"\n❌ Failed to convert {len(failed)} file(s):\n"
            for error in failed[:3]:  # Show first 3 errors
                result_message += f"  • {error}\n"
            if len(failed) > 3:
                result_message += f"  ... and {len(failed) - 3} more errors\n"

        # Update status
        if successful and not failed:
            self.update_status(f"✅ All {len(successful)} files converted successfully!", self.success_color)
        elif successful and failed:
            self.update_status(f"⚠️ {len(successful)} successful, {len(failed)} failed", '#f39c12')
        else:
            self.update_status(f"❌ All conversions failed", self.error_color)

        # Show detailed results
        messagebox.showinfo("Conversion Results", result_message)
        
        # Open output folder if requested and there were successful conversions
        if self.open_output_folder.get() and successful and self.selected_files:
            self.open_folder(os.path.dirname(self.selected_files[0]))

    def open_folder(self, path):
        try:
            if platform.system() == "Windows":
                os.startfile(path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", path])
            else:  # Linux
                subprocess.run(["xdg-open", path])
        except Exception as e:
            print(f"Could not open folder: {e}")

def main():
    root = tk.Tk()
    app = PDFToWordConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()
